"""
agents/resume/converter.py — Convert tailored resume from Markdown to PDF and DOCX.

Primary: pandoc via subprocess (best quality)
Fallback: python-docx for DOCX, reportlab or weasyprint for PDF
"""

from __future__ import annotations

import shutil
import subprocess
import sys
from pathlib import Path
from typing import Any

_REPO_ROOT = Path(__file__).parent.parent.parent
sys.path.insert(0, str(_REPO_ROOT))

from shared.logger import get_logger

log = get_logger("resume")


def convert_to_pdf(md_path: Path, output_path: Path | None = None) -> Path:
    """
    Convert a Markdown file to PDF using pandoc.

    Args:
        md_path: Path to the source .md file.
        output_path: Output path for PDF. Defaults to same dir with .pdf extension.

    Returns:
        Path to the generated PDF file.
    """
    if output_path is None:
        output_path = md_path.with_suffix(".pdf")

    # Try pandoc first
    if shutil.which("pandoc"):
        try:
            result = subprocess.run(
                [
                    "pandoc",
                    str(md_path),
                    "-o", str(output_path),
                    "--pdf-engine=xelatex",
                    "-V", "geometry:margin=1in",
                    "-V", "fontsize=11pt",
                    "-V", "mainfont=Calibri",
                    "--standalone",
                ],
                capture_output=True,
                text=True,
                timeout=60,
            )
            if result.returncode == 0:
                log.info("pdf_converted_pandoc", path=str(output_path))
                return output_path
            else:
                log.warning("pandoc_pdf_failed", stderr=result.stderr[:200])
                # Try without xelatex
                result2 = subprocess.run(
                    ["pandoc", str(md_path), "-o", str(output_path)],
                    capture_output=True, text=True, timeout=60,
                )
                if result2.returncode == 0:
                    return output_path
        except subprocess.TimeoutExpired:
            log.warning("pandoc_timeout")
        except Exception as e:
            log.warning("pandoc_error", error=str(e))

    # Fallback: weasyprint
    try:
        import weasyprint
        import markdown as md_lib
        html_content = md_lib.markdown(md_path.read_text(encoding="utf-8"))
        html_full = f"<html><body style='font-family:Calibri;font-size:11pt;margin:1in'>{html_content}</body></html>"
        weasyprint.HTML(string=html_full).write_pdf(str(output_path))
        log.info("pdf_converted_weasyprint", path=str(output_path))
        return output_path
    except Exception as e:
        log.warning("weasyprint_failed", error=str(e))

    # Last fallback: create a text file noting conversion failed
    output_path.with_suffix(".pdf.txt").write_text(
        f"PDF conversion failed. Install pandoc or weasyprint.\nSource: {md_path}",
        encoding="utf-8",
    )
    log.error("pdf_conversion_failed", md_path=str(md_path))
    raise RuntimeError(f"PDF conversion failed for {md_path}. Install pandoc.")


def convert_to_docx(md_path: Path, output_path: Path | None = None) -> Path:
    """
    Convert a Markdown file to DOCX using pandoc or python-docx fallback.

    Args:
        md_path: Path to the source .md file.
        output_path: Output path for DOCX.

    Returns:
        Path to the generated DOCX file.
    """
    if output_path is None:
        output_path = md_path.with_suffix(".docx")

    # Try pandoc first (best quality)
    if shutil.which("pandoc"):
        try:
            result = subprocess.run(
                ["pandoc", str(md_path), "-o", str(output_path), "--standalone"],
                capture_output=True, text=True, timeout=60,
            )
            if result.returncode == 0:
                log.info("docx_converted_pandoc", path=str(output_path))
                return output_path
        except Exception as e:
            log.warning("pandoc_docx_error", error=str(e))

    # Fallback: python-docx
    try:
        from docx import Document
        from docx.shared import Pt, Inches
        import re

        doc = Document()

        # Set margins
        for section in doc.sections:
            section.top_margin = Inches(1)
            section.bottom_margin = Inches(1)
            section.left_margin = Inches(1)
            section.right_margin = Inches(1)

        md_content = md_path.read_text(encoding="utf-8")

        for line in md_content.split("\n"):
            line = line.rstrip()
            if line.startswith("# "):
                p = doc.add_heading(line[2:], level=1)
            elif line.startswith("## "):
                p = doc.add_heading(line[3:], level=2)
            elif line.startswith("### "):
                p = doc.add_heading(line[4:], level=3)
            elif line.startswith("- ") or line.startswith("* "):
                doc.add_paragraph(line[2:], style="List Bullet")
            elif line.startswith("**") and line.endswith("**"):
                p = doc.add_paragraph()
                run = p.add_run(line.strip("*"))
                run.bold = True
            elif line:
                doc.add_paragraph(line)

        doc.save(str(output_path))
        log.info("docx_converted_python_docx", path=str(output_path))
        return output_path

    except Exception as e:
        log.error("docx_conversion_failed", error=str(e))
        raise RuntimeError(f"DOCX conversion failed: {e}. Install python-docx or pandoc.")


def convert_all(md_path: Path, output_dir: Path | None = None) -> dict[str, Path]:
    """
    Convert a Markdown resume to both PDF and DOCX.

    Args:
        md_path: Source Markdown file.
        output_dir: Output directory. Defaults to same as md_path.

    Returns:
        Dict with 'pdf' and 'docx' keys mapping to output paths.
    """
    if output_dir is None:
        output_dir = md_path.parent
    output_dir.mkdir(parents=True, exist_ok=True)

    stem = md_path.stem
    results: dict[str, Path] = {}

    try:
        pdf_path = convert_to_pdf(md_path, output_dir / f"{stem}.pdf")
        results["pdf"] = pdf_path
    except Exception as e:
        log.error("pdf_failed", error=str(e))

    try:
        docx_path = convert_to_docx(md_path, output_dir / f"{stem}.docx")
        results["docx"] = docx_path
    except Exception as e:
        log.error("docx_failed", error=str(e))

    return results
